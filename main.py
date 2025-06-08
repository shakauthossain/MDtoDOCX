from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from PIL import Image
from io import BytesIO
import markdown2
import io

app = FastAPI()

# === Markdown to HTML Utility Functions ===

def remove_empty_paragraphs_around(soup, tag_names):
    for tag_name in tag_names:
        for tag in soup.find_all(tag_name):
            for prev in tag.find_all_previous():
                if prev.name == "p" and not prev.text.strip():
                    prev.decompose()
                    break
                elif prev.name not in ["p", "br", None]:
                    break
            for next_ in tag.find_all_next():
                if next_.name == "p" and not next_.text.strip():
                    next_.decompose()
                    break
                elif next_.name not in ["p", "br", None]:
                    break

def clean_extra_spacing_around_tables(soup):
    for p in soup.find_all("p"):
        if not p.text.strip():
            p.decompose()

    for table in soup.find_all("table"):
        next_sibling = table.find_next_sibling()
        while next_sibling and (next_sibling.name == "br" or (next_sibling.name == "p" and not next_sibling.text.strip())):
            temp = next_sibling.find_next_sibling()
            next_sibling.decompose()
            next_sibling = temp

def add_table_borders_to_html(html_content: str) -> str:
    soup = BeautifulSoup(html_content, "html.parser")

    for table in soup.find_all("table"):
        table['border'] = "1"
        table['style'] = "border: 1px solid black; border-collapse: collapse; width: 100%;"

        first_row = table.find("tr")
        if first_row:
            col_count = len(first_row.find_all(["td", "th"]))
            colgroup = soup.new_tag("colgroup")
            for _ in range(col_count):
                col = soup.new_tag("col")
                col['style'] = "width: {}%;".format(round(100 / col_count))
                colgroup.append(col)
            table.insert(0, colgroup)

        rows = table.find_all("tr")
        if rows:
            thead = soup.new_tag("thead")
            thead.append(rows[0])
            tbody = soup.new_tag("tbody")
            for row in rows[1:]:
                tbody.append(row)
            table.append(thead)
            table.append(tbody)

        for row in table.find_all("tr"):
            for cell in row.find_all(["th", "td"]):
                existing_style = cell.get('style', '')
                new_style = "border: 1px solid black; padding: 6px;"
                cell['style'] = f"{existing_style} {new_style}".strip()

    return str(soup)

# === API 1: Markdown to HTML ===

@app.post("/convert-md-to-html")
async def convert_md_to_html(request: Request):
    data = await request.json()
    md_text = data.get("markdown", "")
    client_name = data.get("client_name", "Client").strip()

    if not md_text:
        return {"error": "No markdown text provided"}

    html = markdown2.markdown(md_text, extras=[
        "tables",
        "fenced-code-blocks",
        "cuddled-lists",
        "footnotes"
    ])

    soup = BeautifulSoup(html, "html.parser")
    remove_empty_paragraphs_around(soup, ["table", "img", "h1", "h2", "h3", "h4", "h5", "h6"])
    clean_extra_spacing_around_tables(soup)

    cleaned_html = add_table_borders_to_html(str(soup))

    html_bytes = cleaned_html.encode("utf-8")
    html_io = BytesIO(html_bytes)
    html_io.seek(0)

    safe_client_name = "".join(c for c in client_name if c.isalnum() or c in (" ", "_", "-")).strip()
    filename = f"Proposal for {safe_client_name}.html"

    headers = {
        'Content-Disposition': f'attachment; filename="{filename}"'
    }

    return StreamingResponse(
        html_io,
        media_type='text/html',
        headers=headers
    )

# === API 2: Merge Cover Image with DOCX ===

@app.post("/merge-cover-docx")
async def merge_cover_docx(
    cover_image: UploadFile = File(...),
    docx_file: UploadFile = File(...)
):
    try:
        cover_bytes = await cover_image.read()
        docx_bytes = await docx_file.read()

        # Create DOCX and insert image
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        width = section.page_width - section.left_margin - section.right_margin

        doc.add_picture(BytesIO(cover_bytes), width=width)
        doc.add_page_break()

        # Load and append content
        proposal = Document(BytesIO(docx_bytes))
        for elem in proposal.element.body:
            doc.element.body.append(elem)

        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        return StreamingResponse(
            output_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="Merged_Proposal_With_Cover.docx"'}
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
